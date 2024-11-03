import docx
from docx import Document


def read_docx(file_path, idx):
    document = Document(file_path)
    context = ""
    for para in document.paragraphs[idx]:
        context += para.text + '\n'
    return context


def find_file_length(file_path):
    document = Document(file_path)
    return len(document.paragraphs)


if __name__ == '__main__':
    path = r'xxx\光明街道年鉴2022.docx'
    idx = slice(2000, 2080, None)
    context = read_docx(path, idx)
    # find out how many tokens are there in the document
    print(context)
    print(len(context))
